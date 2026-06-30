from docx import Document

from docx.enum.text import (
    WD_COLOR_INDEX
)

# ==========================================
# COLOR MAPPER
# ==========================================

def get_docx_color(color):

    if color == "red":

        return WD_COLOR_INDEX.RED

    elif color == "orange":

        return WD_COLOR_INDEX.YELLOW

    elif color == "yellow":

        return WD_COLOR_INDEX.BRIGHT_GREEN

    return WD_COLOR_INDEX.GRAY_25

# ==========================================
# GENERATE HIGHLIGHTED DOCX
# ==========================================

def generate_highlighted_docx(

    original_text,

    analyzed_sentences,

    output_path
):

    doc = Document()

    heading = doc.add_heading(
        "DeepFakeConnect Forensic Report",
        level=1
    )

    heading.runs[0].bold = True

    p = doc.add_paragraph()

    # ======================================
    # SENTENCE COLORING
    # ======================================

    for item in analyzed_sentences:

        sentence = item["sentence"]

        color = item["color"]

        score = item["score"]

        risk = item["risk"]

        run = p.add_run(
            sentence + " "
        )

        run.font.highlight_color = (
            get_docx_color(color)
        )

        # Optional forensic annotation

        annotation = p.add_run(

            f"[AI Score: {score}% | {risk}] "

        )

        annotation.bold = True

    # ======================================
    # SAVE FILE
    # ======================================

    doc.save(output_path)

    return output_path